from docx import Document
from docx import Document
import os
from striprtf.striprtf import rtf_to_text

def extract_text_from_file(file_path):
    """Извлекает текст из файлов разных форматов"""
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    try:
        if ext == '.docx':
            # Для DOCX файлов
            doc = Document(file_path)
            return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            
        elif ext == '.rtf':
            # Для RTF файлов
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
                return rtf_to_text(rtf_content)
                
        elif ext == '.txt':
            # Для TXT файлов
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        else:
            # Попробуем как обычный текст
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
    except Exception as e:
        print(f"Ошибка при обработке файла {file_path}: {e}")
        return ""
    
def extract_all_text(docx_path):
    doc = Document(docx_path)
    text = []
    
    # Параграфы
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
    
    # Текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)
    
    return '\n'.join(text)

#text = extract_all_text("data/resume_1_IT.rtf.docx")
#print("Извлеченный текст:")
#print(repr(text))  # repr покажет специальные символы
#print("\n---\n")
#print(text)